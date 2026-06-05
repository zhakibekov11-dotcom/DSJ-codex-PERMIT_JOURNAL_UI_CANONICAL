#!/usr/bin/env python3

import json
import sys
from pathlib import Path

from docx import Document


def read_json_stdin():
    return json.loads(sys.stdin.buffer.read().decode("utf-8"))


def remove_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def remove_table(table):
    element = table._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def clone_run_style(source_run, target_run):
    if source_run is None:
        return

    target_run.style = source_run.style
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline

    source_font = source_run.font
    target_font = target_run.font
    if source_font.name:
        target_font.name = source_font.name
    if source_font.size:
        target_font.size = source_font.size
    if source_font.color and source_font.color.rgb:
        target_font.color.rgb = source_font.color.rgb


def set_paragraph_text(paragraph, text):
    source_run = paragraph.runs[0] if paragraph.runs else None

    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)

    parts = text.split("\n")
    for index, part in enumerate(parts):
        if index > 0:
            line_break = paragraph.add_run()
            clone_run_style(source_run, line_break)
            line_break.add_break()

        run = paragraph.add_run(part)
        clone_run_style(source_run, run)


def find_paragraph_index(paragraphs, marker):
    for index, paragraph in enumerate(paragraphs):
        if marker in paragraph.text:
            return index
    raise RuntimeError(f"Не найден маркер шаблона: {marker}")


def build_recipient_block(recipients):
    blocks = []
    for recipient in recipients:
        lines = [recipient["companyName"].strip()]
        contact_line = ", ".join(
            value
            for value in [recipient.get("contactName", "").strip(), recipient.get("contactPosition", "").strip()]
            if value
        )
        if contact_line:
            lines.append(f"Вниманию: {contact_line}")
        if recipient.get("contactEmail"):
            lines.append(recipient["contactEmail"].strip())
        blocks.append("\n".join(lines))
    return "\n\n".join(blocks) if blocks else "Получатель не указан"


def main():
    if len(sys.argv) != 3:
        print("Usage: generate_correspondence_docx.py <template.docx> <output.docx>", file=sys.stderr)
        sys.exit(1)

    template_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    payload = read_json_stdin()

    document = Document(template_path)
    paragraphs = list(document.paragraphs)

    registry_index = find_paragraph_index(paragraphs, "исх: 0025SC-2025")
    recipient_index = find_paragraph_index(paragraphs, "ТОО “Оркен”")
    heading_index = find_paragraph_index(paragraphs, "Коммерческое предложение")
    body_index = find_paragraph_index(paragraphs, "Разработка проекта плана ликвидации последствий недропользования")
    signoff_index = find_paragraph_index(paragraphs, "С уважением,")
    signature_index = find_paragraph_index(paragraphs, "Директор ТОО «Stroy Company 2030»")

    registry_paragraph = paragraphs[registry_index]
    recipient_paragraph = paragraphs[recipient_index]
    heading_paragraph = paragraphs[heading_index]
    body_paragraph = paragraphs[body_index]

    for paragraph in paragraphs[body_index + 1 : signoff_index]:
        remove_paragraph(paragraph)

    for paragraph in paragraphs[signature_index + 1 :]:
        remove_paragraph(paragraph)

    for table in list(document.tables):
        remove_table(table)

    set_paragraph_text(
        registry_paragraph,
        f"исх: {payload['registryNumber']} от {payload['issueDateRu']}",
    )
    set_paragraph_text(recipient_paragraph, build_recipient_block(payload["recipients"]))
    set_paragraph_text(heading_paragraph, payload["heading"])
    set_paragraph_text(body_paragraph, payload["body"])

    document.save(output_path)


if __name__ == "__main__":
    main()
