#!/usr/bin/env python3

from __future__ import annotations

import json
import re
import sys
from copy import deepcopy
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


PLACEHOLDER_RE = re.compile(r"\{\{[A-Z0-9_]+\}\}")
REQUIRED_FIELDS = {
    "{{KB_NUMBER}}",
    "{{REGISTRATION_NUMBER}}",
    "{{FULL_NAME_RU}}",
    "{{FULL_NAME_KZ}}",
    "{{PROFESSION_RU}}",
    "{{PROFESSION_KZ}}",
    "{{PROTOCOL_NUMBER_DISPLAY}}",
}
FONT_SIZE_OVERRIDES_PT = {
    "{{KB_NUMBER}}": 11,
}
APPEND_LINE_BREAK_AFTER = {
    "{{KB_NUMBER}}",
}


def read_json_stdin() -> object:
    return json.loads(sys.stdin.buffer.read().decode("utf-8"))


def normalize_value(value: object) -> str:
    if value is None:
        return ""
    return str(value).replace("\r\n", " ").replace("\n", " ").strip()


def normalize_row(row: object) -> dict[str, str]:
    if not isinstance(row, dict):
        raise RuntimeError("Одна из строк свидетельства передана в неверном формате.")

    fields = row.get("fields") if "fields" in row else row
    if not isinstance(fields, dict):
        raise RuntimeError("Одна из строк свидетельства не содержит fields.")

    normalized = {key: normalize_value(value) for key, value in fields.items()}
    for field in REQUIRED_FIELDS:
        if not normalized.get(field):
            raise RuntimeError(f"Для свидетельства не заполнено поле {field}.")

    return normalized


def validate_payload(payload: object) -> list[dict[str, str]]:
    if not isinstance(payload, dict):
        raise RuntimeError("Передан некорректный payload свидетельства ПС.")

    rows = payload.get("rows")
    if not isinstance(rows, list) or not rows:
        raise RuntimeError("Для свидетельства ПС нужен хотя бы один получатель.")

    return [normalize_row(row) for row in rows]


def replace_placeholder_runs(paragraph, fields: dict[str, str]) -> None:
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue

        replaced = text
        font_size_override = None
        append_line_break = False
        for placeholder, value in fields.items():
            if placeholder in replaced:
                replaced = replaced.replace(placeholder, value)
                font_size_override = FONT_SIZE_OVERRIDES_PT.get(placeholder, font_size_override)
                if placeholder in APPEND_LINE_BREAK_AFTER:
                    append_line_break = True

        run.text = replaced
        if font_size_override is not None:
            run.font.name = "Times New Roman"
            run.font.size = Pt(font_size_override)
            r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
            r_fonts.set(qn("w:ascii"), "Times New Roman")
            r_fonts.set(qn("w:hAnsi"), "Times New Roman")
            r_fonts.set(qn("w:eastAsia"), "Times New Roman")
            r_fonts.set(qn("w:cs"), "Times New Roman")
        if append_line_break:
            run.add_break()

    paragraph_text = "".join(run.text for run in paragraph.runs)
    leftovers = PLACEHOLDER_RE.findall(paragraph_text)
    if leftovers:
        raise RuntimeError(f"В шаблоне остались необработанные коды: {', '.join(sorted(set(leftovers)))}")


def walk_paragraphs(parent):
    for paragraph in parent.paragraphs:
        yield paragraph

    for table in parent.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from walk_paragraphs(cell)


def fill_document(document: Document, row: dict[str, str]) -> None:
    for paragraph in walk_paragraphs(document):
        replace_placeholder_runs(paragraph, row)


def render_document_bytes(template_path: Path, row: dict[str, str]) -> bytes:
    document = Document(template_path)
    fill_document(document, row)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_preview_document(template_path: Path, output_path: Path, rows: list[dict[str, str]]) -> None:
    rendered_docs = [Document(BytesIO(render_document_bytes(template_path, row))) for row in rows]
    base_document = rendered_docs[0]

    for document in rendered_docs[1:]:
        base_document.add_page_break()
        for child in list(document.element.body):
            if child.tag == qn("w:sectPr"):
                continue
            base_document.element.body.append(deepcopy(child))

    base_document.save(output_path)


def main() -> int:
    if len(sys.argv) != 3:
        print(
            "Usage: generate_ps_witness_certificate.py <template.docx> <output.docx>",
            file=sys.stderr,
        )
        return 1

    template_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])

    if not template_path.exists():
        print("Шаблон свидетельства ПС не найден.", file=sys.stderr)
        return 1

    try:
        rows = validate_payload(read_json_stdin())
        build_preview_document(template_path, output_path, rows)
    except Exception as error:  # noqa: BLE001
        print(str(error), file=sys.stderr)
        return 1

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
