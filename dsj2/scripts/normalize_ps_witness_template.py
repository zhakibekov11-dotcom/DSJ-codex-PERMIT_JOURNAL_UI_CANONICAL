#!/usr/bin/env python3

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt


def clear_paragraph(paragraph) -> None:
    paragraph_element = paragraph._p
    for child in list(paragraph_element):
        if child.tag.endswith("}pPr"):
            continue
        paragraph_element.remove(child)


def set_paragraph_text(paragraph, text: str, *, bold: bool | None = None, size_pt: float | None = None):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    return run


def set_cell_paragraph_text(
    cell,
    paragraph_index: int,
    text: str,
    *,
    bold: bool | None = None,
    size_pt: float | None = None,
) -> None:
    paragraph = cell.paragraphs[paragraph_index]
    set_paragraph_text(paragraph, text, bold=bold, size_pt=size_pt)


def clear_other_paragraphs(cell, keep_index: int = 0) -> None:
    for index, paragraph in enumerate(cell.paragraphs):
        if index == keep_index:
            continue
        set_paragraph_text(paragraph, "")


def unique_cells(row):
    seen = set()
    result = []
    for cell in row.cells:
        identifier = id(cell._tc)
        if identifier in seen:
            continue
        seen.add(identifier)
        result.append(cell)
    return result


def normalize_template(source_path: Path, output_path: Path) -> None:
    document = Document(source_path)
    main_table = document.tables[1]
    kz_cell = main_table.rows[0].cells[0]
    ru_cell = main_table.rows[0].cells[1]

    set_cell_paragraph_text(kz_cell, 3, "{{KB_NUMBER}}", bold=True, size_pt=16)
    set_cell_paragraph_text(ru_cell, 3, "{{KB_NUMBER}}", bold=True, size_pt=16)
    set_cell_paragraph_text(
        kz_cell,
        16,
        "Тіркеу нөмері № {{REGISTRATION_NUMBER}}",
        bold=True,
        size_pt=8,
    )
    set_cell_paragraph_text(
        ru_cell,
        14,
        "Регистрационный номер № {{REGISTRATION_NUMBER}}",
        bold=True,
        size_pt=8,
    )

    kz_tables = kz_cell.tables
    ru_tables = ru_cell.tables

    # KZ side
    top_kz = kz_tables[0]
    for cell in unique_cells(top_kz.rows[0]):
        set_cell_paragraph_text(cell, 0, "Осы куәлік: {{FULL_NAME_KZ}}", size_pt=9)
    set_cell_paragraph_text(top_kz.rows[3].cells[1], 0, "{{TRAINING_START_YEAR_FULL}}", size_pt=9)
    set_cell_paragraph_text(top_kz.rows[3].cells[5], 0, "{{TRAINING_START_DAY}}", size_pt=9)
    set_cell_paragraph_text(top_kz.rows[3].cells[7], 0, "{{TRAINING_START_MONTH_KZ}}", size_pt=9)
    set_cell_paragraph_text(top_kz.rows[3].cells[9], 0, "{{TRAINING_END_YEAR_FULL}}", size_pt=9)
    set_cell_paragraph_text(top_kz.rows[3].cells[12], 0, "{{TRAINING_END_DAY}}", size_pt=9)
    set_cell_paragraph_text(top_kz.rows[3].cells[14], 0, "{{TRAINING_END_MONTH_KZ}}", size_pt=9)

    org_kz = kz_tables[1]
    set_cell_paragraph_text(org_kz.rows[0].cells[2], 0, "{{EDU_ORG_KZ}}", size_pt=8)
    set_cell_paragraph_text(org_kz.rows[0].cells[3], 0, "{{EDU_ORG_KZ}}", size_pt=8)
    set_cell_paragraph_text(org_kz.rows[2].cells[0], 0, "{{TRAINING_END_YEAR_FULL}}", size_pt=9)
    set_cell_paragraph_text(org_kz.rows[2].cells[2], 0, "{{PROFESSION_KZ}}", bold=True, size_pt=8)
    set_cell_paragraph_text(org_kz.rows[2].cells[3], 0, "{{PROFESSION_KZ}}", bold=True, size_pt=8)

    commission_kz = kz_tables[3]
    set_cell_paragraph_text(
        commission_kz.rows[0].cells[2],
        0,
        "{{TRAINING_END_YEAR_FULL}}",
        size_pt=8,
    )
    set_cell_paragraph_text(
        commission_kz.rows[0].cells[3],
        0,
        "{{TRAINING_END_YEAR_FULL}}",
        size_pt=8,
    )
    set_cell_paragraph_text(
        commission_kz.rows[0].cells[6],
        0,
        "{{TRAINING_END_MONTH_KZ}}",
        size_pt=8,
    )
    set_cell_paragraph_text(
        commission_kz.rows[0].cells[7],
        0,
        "{{TRAINING_END_MONTH_KZ}}",
        size_pt=8,
    )
    set_cell_paragraph_text(
        commission_kz.rows[0].cells[8],
        0,
        "{{TRAINING_END_MONTH_KZ}}",
        size_pt=8,
    )
    set_cell_paragraph_text(
        commission_kz.rows[1].cells[1],
        0,
        "{{PROTOCOL_NUMBER_DISPLAY}}",
        size_pt=8,
    )
    set_cell_paragraph_text(
        commission_kz.rows[1].cells[2],
        0,
        "{{PROTOCOL_NUMBER_DISPLAY}}",
        size_pt=8,
    )
    for cell in unique_cells(commission_kz.rows[2]):
        set_cell_paragraph_text(cell, 0, "{{PROFESSION_KZ}}", bold=True, size_pt=8)

    issue_kz = kz_tables[4]
    set_cell_paragraph_text(issue_kz.rows[0].cells[1], 0, "{{ISSUE_DAY}}", size_pt=8)
    set_cell_paragraph_text(issue_kz.rows[0].cells[3], 0, "{{ISSUE_MONTH_KZ}}", size_pt=8)
    set_cell_paragraph_text(issue_kz.rows[0].cells[5], 0, "{{ISSUE_YEAR_SHORT}}", size_pt=8)

    # RU side
    top_ru = ru_tables[0]
    set_cell_paragraph_text(top_ru.rows[0].cells[0], 0, "Настоящее свидетельство выдано: {{FULL_NAME_RU}}", size_pt=9)

    training_ru = ru_tables[1]
    set_cell_paragraph_text(training_ru.rows[0].cells[2], 0, "{{TRAINING_START_DAY}}", size_pt=9)
    set_cell_paragraph_text(training_ru.rows[0].cells[4], 0, "{{TRAINING_START_MONTH_RU}}", size_pt=9)
    set_cell_paragraph_text(training_ru.rows[0].cells[6], 0, "{{TRAINING_START_YEAR_SHORT}}", size_pt=9)
    set_cell_paragraph_text(training_ru.rows[0].cells[10], 0, "{{TRAINING_END_DAY}}", size_pt=9)
    set_cell_paragraph_text(training_ru.rows[0].cells[12], 0, "{{TRAINING_END_MONTH_RU}}", size_pt=9)
    set_cell_paragraph_text(training_ru.rows[0].cells[14], 0, "{{TRAINING_END_YEAR_SHORT}}", size_pt=9)

    org_ru = ru_tables[2]
    for cell in unique_cells(org_ru.rows[0]):
        set_cell_paragraph_text(cell, 0, "{{EDU_ORG_RU}}", size_pt=8)
    set_cell_paragraph_text(org_ru.rows[2].cells[1], 0, "{{TRAINING_END_YEAR_FULL}}", size_pt=9)
    set_cell_paragraph_text(org_ru.rows[2].cells[2], 0, "{{TRAINING_END_YEAR_FULL}}", size_pt=9)
    set_cell_paragraph_text(org_ru.rows[3].cells[3], 0, "{{PROFESSION_RU}}", bold=True, size_pt=9)
    set_cell_paragraph_text(org_ru.rows[3].cells[4], 0, "{{PROFESSION_RU}}", bold=True, size_pt=9)
    set_cell_paragraph_text(org_ru.rows[3].cells[5], 0, "{{PROFESSION_RU}}", bold=True, size_pt=9)

    commission_ru = ru_tables[4]
    sentence = (
        "Решением квалификационной комиссии от "
        "«{{ISSUE_DAY}}» {{ISSUE_MONTH_RU}} {{TRAINING_END_YEAR_FULL}} г. "
        "№ протокола {{PROTOCOL_NUMBER_DISPLAY}} ему (ей) присвоена квалификация {{PROFESSION_RU}}"
    )
    for cell in unique_cells(commission_ru.rows[0]):
        set_cell_paragraph_text(cell, 0, sentence, size_pt=8)
        clear_other_paragraphs(cell, keep_index=0)

    issue_ru = ru_tables[5]
    set_cell_paragraph_text(issue_ru.rows[0].cells[1], 0, "{{ISSUE_DAY}}", size_pt=8)
    set_cell_paragraph_text(issue_ru.rows[0].cells[3], 0, "{{ISSUE_MONTH_RU}}", size_pt=8)
    set_cell_paragraph_text(issue_ru.rows[0].cells[5], 0, "{{ISSUE_YEAR_SHORT}}", size_pt=8)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> int:
    if len(sys.argv) != 3:
        print(
            "Usage: normalize_ps_witness_template.py <source.docx> <output.docx>",
            file=sys.stderr,
        )
        return 1

    source_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])

    if not source_path.exists():
        print("Исходный шаблон свидетельства ПС не найден.", file=sys.stderr)
        return 1

    normalize_template(source_path, output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
