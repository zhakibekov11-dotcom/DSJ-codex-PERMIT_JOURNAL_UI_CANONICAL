#!/usr/bin/env python3

import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION


CATEGORY_LABELS = {
    "LOCAL_ACT": "Локальные акты",
    "ORDER": "Приказы",
    "INSTRUCTION": "Инструкции",
    "JOURNAL": "Журналы",
    "TRAINING_CERTIFICATION": "Обучение и аттестация",
}

STATUS_LABELS = {
    "DRAFT": "Черновик",
    "ACTIVE": "Действует",
    "ARCHIVED": "В архиве",
}


def format_ru_date(value):
    if not value:
        return "Не указана"

    try:
        date = datetime.fromisoformat(value.replace("Z", "+00:00"))
    except ValueError:
        return value

    return date.strftime("%d.%m.%Y")


def add_meta_line(document, label, value):
    paragraph = document.add_paragraph()
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    paragraph.add_run(value or "Не указано")


def main():
    if len(sys.argv) != 2:
        print(
            "Usage: generate_company_document_docx.py <output.docx>",
            file=sys.stderr,
        )
        sys.exit(1)

    output_path = Path(sys.argv[1])
    payload = json.load(sys.stdin)

    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.top_margin = 720000
    section.bottom_margin = 720000
    section.left_margin = 720000
    section.right_margin = 720000

    heading = document.add_heading(payload["title"].strip(), level=0)
    heading.alignment = 0

    subtitle = document.add_paragraph()
    subtitle_run = subtitle.add_run(
        f"{CATEGORY_LABELS.get(payload['category'], payload['category'])} | {payload['documentName'].strip()}"
    )
    subtitle_run.italic = True

    add_meta_line(document, "Компания", payload["companyName"].strip())
    add_meta_line(
        document,
        "Категория",
        CATEGORY_LABELS.get(payload["category"], payload["category"]),
    )
    add_meta_line(document, "Вид документа", payload["documentName"].strip())
    add_meta_line(
        document,
        "Статус",
        STATUS_LABELS.get(payload["status"], payload["status"]),
    )
    add_meta_line(document, "Дата документа", format_ru_date(payload.get("issueDate")))
    add_meta_line(document, "Подготовил", payload["createdByUserName"].strip())
    add_meta_line(document, "Обновлено", format_ru_date(payload.get("updatedAt")))

    if payload.get("summary"):
        document.add_heading("Краткое описание", level=1)
        document.add_paragraph(payload["summary"].strip())

    document.add_heading("Содержание", level=1)
    paragraphs = [
        paragraph.strip()
        for paragraph in payload["body"].split("\n\n")
        if paragraph.strip()
    ]

    if not paragraphs:
        document.add_paragraph("Содержание не заполнено.")
    else:
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)

    document.save(output_path)


if __name__ == "__main__":
    main()
